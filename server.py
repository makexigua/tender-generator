# -*- coding: utf-8 -*-
import asyncio
import logging
import os
import re
from contextlib import asynccontextmanager
from functools import lru_cache
from pathlib import Path
from urllib.parse import quote

import paramiko
from pymysql import connect
from mcp.server import Server
from mcp.server.streamable_http_manager import StreamableHTTPSessionManager
from mcp.types import Resource, TextContent, Tool
from starlette.applications import Starlette
from starlette.responses import JSONResponse
from starlette.routing import Route
import uvicorn

from docx import Document
import uuid

from upd_status import update_biding_doc_status

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("mysql_mcp_server")


def load_env_file():
    """
    读取项目根目录下的 .env。
    大白话：MCP server 也要知道回写接口地址，所以这里和主流程一样补一份轻量的 .env 读取。
    """
    env_path = Path(__file__).resolve().parent / ".env"
    if not env_path.exists():
        return

    for raw_line in env_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip()
        if len(value) >= 2 and value[0] == value[-1] and value[0] in {"'", '"'}:
            value = value[1:-1]
        os.environ.setdefault(key, value)


def get_required_env(name: str) -> str:
    value = os.getenv(name, "").strip()
    if not value:
        raise ValueError(f"缺少配置：{name}")
    return value


def get_public_download_base_url() -> str:
    """
    获取对外下载地址前缀。
    默认值和 main_async_threaded.py 里发给钉钉的附件下载域名保持一致。
    """
    return os.getenv(
        "PUBLIC_DOWNLOAD_BASE_URL",
        "https://ai-assistant.4-xiang.com/download",
    ).rstrip("/")


def get_generated_docs_dir() -> Path:
    """
    主流程和 MCP 服务共用文档目录，避免“主流程能写、下载路由找不到”的问题。

    优先级：
    1) GENERATED_DOCS_DIR
    2) 兼容旧变量 ATTACHMENT_LOCAL_DIR
    3) root 用户默认 /root/generated_docs
    4) 非 root 用户默认项目内「招标文件」目录
    """
    env_dir = os.getenv("GENERATED_DOCS_DIR", "").strip()
    if not env_dir:
        env_dir = os.getenv("ATTACHMENT_LOCAL_DIR", "").strip()
    if env_dir:
        return Path(env_dir)

    try:
        if hasattr(os, "geteuid") and os.geteuid() == 0:
            return Path("/root/generated_docs")
    except Exception:
        pass

    return Path(__file__).resolve().parent / "招标文件"


def update_tender_generation_status(
    uid: str,
    tender_uid: str,
    status: int,
    memo: str,
    bid_doc_location_url: str = "",
) -> dict:
    """
    回写投标文件生成状态。
    status=3 表示生成成功，status=4 表示生成失败。
    """
    api_url = get_required_env("https://api.4-xiang.com/admin/tender/biding_doc/upd_status")
    return update_biding_doc_status(
        api_url=api_url,
        uid=uid,
        tender_uid=tender_uid,
        bid_doc_location_url=bid_doc_location_url,
        status=status,
        memo=memo,
    )


load_env_file()

# ================== 数据库配置 ==================
def get_db_config():
    return {
        "host": os.getenv("MYSQL_HOST", "***"),
        "port": int(os.getenv("MYSQL_PORT", "3306")),
        "user": os.getenv("MYSQL_USER", "prod"),
        "password": os.getenv("MYSQL_PASSWORD", "****"),
        "database": os.getenv("MYSQL_DATABASE", "***"),
        "charset": "utf8mb4",
        "autocommit": True,
    }


# ================== 分表路由 ==================
def extract_park_name(query: str):
    match = re.search(r"park\.name\s*=\s*['\"](.+?)['\"]", query, re.IGNORECASE)
    return match.group(1) if match else None


@lru_cache(maxsize=1000)
def get_sharding_num_cached(park_name: str):
    config = get_db_config()
    with connect(**config) as conn:
        with conn.cursor() as cursor:
            cursor.execute("SELECT sharding_num FROM park WHERE name=%s LIMIT 1", (park_name,))
            row = cursor.fetchone()
            return row[0] if row else None


def rewrite_sql(query: str) -> str:
    if "toll_detail_final" not in query.lower():
        return query

    park_name = extract_park_name(query)
    if not park_name:
        raise ValueError("查询分表必须带 park.name 条件")

    # 查询park表获取sharding_num和uid
    config = get_db_config()
    with connect(**config) as conn:
        with conn.cursor() as cursor:
            cursor.execute("SELECT sharding_num, uid FROM park WHERE name=%s LIMIT 1", (park_name,))
            row = cursor.fetchone()
            if not row:
                raise ValueError(f"未找到车场: {park_name}")
            sharding_num, park_uid = row

    real_table = f"toll_detail_final_{sharding_num}"

    # 替换表名
    query = re.sub(r"\btoll_detail_final\b", real_table, query, flags=re.IGNORECASE)

    # 关键：将 park.name = 'xxx' 或 park.name = "xxx" 替换为 park_uid = 'uid'
    query = re.sub(
        r"park\.name\s*=\s*['\"]([^'\"]+)['\"]",
        f"park_uid = '{park_uid}'",
        query,
        flags=re.IGNORECASE
    )

    return query


# ================== SQL安全 ==================
def validate_sql(query: str):
    forbidden = ["drop", "delete", "truncate", "alter"]
    lower_q = query.lower()
    for word in forbidden:
        if word in lower_q:
            raise ValueError(f"禁止危险SQL: {word}")


def add_limit(query: str) -> str:
    if "limit" not in query.lower():
        query += " LIMIT 1000"
    return query



# ================== Linux命令执行相关配置和函数 ==================
import subprocess
import datetime
import re

LOG_DIR = "/opt/lcv2/be/logs"
REMOTE_SERVER = "root@120.27.10.214"


# --------------------------
# 执行远程命令
# --------------------------
def run_ssh_command(inner_cmd: str) -> str:
    if "ssh " in inner_cmd:
        raise Exception("禁止嵌套SSH命令")
    
    command = f"ssh {REMOTE_SERVER} \"{inner_cmd}\""
    print(f"【DEBUG】执行命令: {command}")

    try:
        result = subprocess.run(
            command,
            shell=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            timeout=60
        )
        return result.stdout.strip()
    except Exception as e:
        return f"执行失败: {str(e)}"


# --------------------------
# 构造 grep 命令
# --------------------------
def build_grep_cmd(plate: str, log_path: str, filter_cmd: str) -> str:
    cmd = f"grep -a -A 30 -B 30 '{plate}' {log_path}"
    if filter_cmd:
        cmd += f" | {filter_cmd}"
    return cmd


# --------------------------
# 查询 app.log
# --------------------------
def query_app_log(plate: str, filter_cmd: str) -> str:
    log_path = f"{LOG_DIR}/app.log"
    cmd = build_grep_cmd(plate, log_path, filter_cmd)
    return run_ssh_command(cmd)


# --------------------------
# 查询历史日志（8小时回溯）
# --------------------------
def query_history_logs(plate: str, filter_cmd: str) -> str:
    now = datetime.datetime.now()

    # 回溯8小时
    for i in range(1, 9):
        t = now - datetime.timedelta(hours=i)
        date_str = t.strftime("%Y-%m-%d")

        log_path = f"{LOG_DIR}/app.{date_str}.log-*.zip"

        cmd = f"zcat {log_path} | grep -a -A 30 -B 30 '{plate}'"
        if filter_cmd:
            cmd += f" | {filter_cmd}"

        result = run_ssh_command(cmd)

        if result.strip():
            print(f"【DEBUG】命中历史日志: {date_str}")
            return result

    return ""


def parse_time(user_text: str):
    now = datetime.datetime.now()

    # 绝对时间：2026-04-24 00:41:25
    match = re.search(r"\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}", user_text)
    if match:
        return datetime.datetime.strptime(match.group(), "%Y-%m-%d %H:%M:%S")

    # 相对时间
    if "小时前" in user_text:
        hours = int(re.search(r"(\d+)小时", user_text).group(1))
        return now - datetime.timedelta(hours=hours)

    if "分钟前" in user_text:
        mins = int(re.search(r"(\d+)分钟", user_text).group(1))
        return now - datetime.timedelta(minutes=mins)

    if "天前" in user_text:
        days = int(re.search(r"(\d+)天", user_text).group(1))
        return now - datetime.timedelta(days=days)

    return None

def list_remote_logs():
    cmd = f"ls -lt --time-style=full-iso {LOG_DIR}"
    output = run_ssh_command(cmd)

    logs = []

    for line in output.split("\n"):
        parts = line.split()

        if len(parts) < 9:
            continue

        filename = parts[-1]

        # ✅ 只保留 app.log 和历史日志
        if not (filename == "app.log" or re.match(r"app\.\d{4}-\d{2}-\d{2}\.log-\d+\.zip", filename)):
            continue

        date_str = parts[5]
        time_str = parts[6].split(".")[0]  # 去掉毫秒

        try:
            dt = datetime.datetime.strptime(
                f"{date_str} {time_str}",
                "%Y-%m-%d %H:%M:%S"
            )
            logs.append((filename, dt))
        except Exception as e:
            print(f"[解析失败] {line}")
            continue

    print(f"【DEBUG】解析到日志数量: {len(logs)}")

    # 🔥 制打印前几个，方便排查
    for f, t in logs[:5]:
        print(f"[LOG] {f} -> {t}")

    return logs
    
def find_closest_log(target_time, logs):
    print(f"【进入了find_closest_log】函数，target_time: {target_time}")

    if not logs:
        return None

    print(f"*** logs 不为空，数量: {len(logs)} ***")

    # 1️⃣ 找所有 >= target_time 的日志
    future_logs = [log for log in logs if log[1] >= target_time]

    if future_logs:
        # 2️⃣ 取"最接近未来"的那个（最小的）
        chosen = min(future_logs, key=lambda x: x[1])
        print(f"【DEBUG】命中未来最近日志: {chosen[0]} 时间: {chosen[1]}")
        return chosen

    # 3️⃣ 如果没有未来日志 → 取最近过去的（最大的）
    past_logs = [log for log in logs if log[1] < target_time]

    if past_logs:
        chosen = max(past_logs, key=lambda x: x[1])
        print(f"【DEBUG】没有未来日志，取最近过去日志: {chosen[0]} 时间: {chosen[1]}")
        return chosen

    print("【ERROR】没有任何可用日志")
    return None
def query_specific_log(filename, plate, filter_cmd):
    path = f"{LOG_DIR}/{filename}"

    # 🔥 更通用判断
    if filename.endswith(".zip"):
        cmd = f"zcat {path} | grep -a -A 30 -B 30 '{plate}'"
    else:
        cmd = f"grep -a -A 30 -B 30 '{plate}' {path}"

    if filter_cmd:
        cmd += f" | {filter_cmd}"

    print(f"【DEBUG】最终执行命令: {cmd}")

    return run_ssh_command(cmd)
                
# --------------------------
# 1. 关键词 → 二次过滤 映射
# --------------------------
FILTER_MAPPING = [
    {
        "keywords": ["抬杆", "发送指令", "开闸", "抬杆指令"],
        "filter": "grep -i '开闸'",
        "desc": "查询抬杆指令"
    },
    {
        "keywords": ["白名单", "月租", "长租", "VIP"],
        "filter": "grep -i '白名单'",
        "desc": "查询白名单状态"
    },
    {
        "keywords": ["欠费", "未缴费", "逃费"],
        "filter": "grep -i '欠费'",
        "desc": "查询欠费记录"
    }
]

# --------------------------
# 2. 从用户问题提取车牌
# --------------------------
def extract_license_plate(text: str) -> str:
    pattern = r"[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤青藏川宁琼][A-Za-z0-9]{5,8}"
    match = re.search(pattern, text)
    return match.group().strip() if match else ""

# --------------------------
# 3. 匹配用户意图，返回二次过滤命令
# --------------------------
def get_filter_command(user_text: str) -> str:
    for item in FILTER_MAPPING:
        for kw in item["keywords"]:
            if kw in user_text:
                return item["filter"]
    return ""

# ============================================================
# 新增：查询出场车道 / 查询车场名称 功能
# ============================================================

def is_park_lane_query(user_text: str) -> bool:
    """判断用户是否想查询出场车道或车场名称"""
    keywords = ["查询出场车道", "查询车场名称"]
    return any(kw in user_text for kw in keywords)


def query_park_lane_info(plate: str) -> str:
    """
    执行特定命令查询车场名称和车道信息
    命令格式: ssh root@120.27.10.214 "grep '车牌' /opt/lcv2/be/logs/app.log|grep -i '最近的过车记录的车道是在'"
    """
    inner_cmd = f"grep '{plate}' /opt/lcv2/be/logs/app.log | grep -i '最近的过车记录的车道是在'"
    return run_ssh_command(inner_cmd)


def parse_park_lane_output(output: str) -> dict:
    """
    解析日志输出，提取：
    - 车场名称（@LOG@< 和 >车辆 之间）
    - 入口车道（"最近的过车记录的车道是在:" 后面，" ，" 前面）
    - 出口车道（"本次车道:" 后面）
    
    示例输入:
    2026-04-28 20:03:08,130 [camera-mqtt-23] INFO  [c.s.h.l.c.i.e.CrossCaseOfInExchanger.?:?]: @LOG@<湘府文化公园>车辆：豫NW787X 最近的过车记录的车道是在:北入口 ，本次车道:南出口
    """
    result = {
        "park_name": None,
        "entry_lane": None,
        "exit_lane": None,
        "raw": output
    }
    
    if not output or "执行失败" in output:
        return result
    
    # 提取车场名称: @LOG@<XXXX>车辆
    park_match = re.search(r'@LOG@<(.+?)>车辆', output)
    if park_match:
        result["park_name"] = park_match.group(1)
    
    # 提取入口车道: "最近的过车记录的车道是在:" 后面到 " ，" 之前
    entry_match = re.search(r'最近的过车记录的车道是在[:：](.+?)\s*，', output)
    if entry_match:
        result["entry_lane"] = entry_match.group(1).strip()
    
    # 提取出口车道: "本次车道:" 后面
    exit_match = re.search(r'本次车道[:：](.+?)(?:\s|$)', output)
    if exit_match:
        result["exit_lane"] = exit_match.group(1).strip()
    
    return result


def format_park_lane_response(parsed: dict, plate: str) -> str:
    """格式化输出查询结果"""
    if not parsed["park_name"]:
        return f"未查询到车牌 {plate} 的车场和车道信息，原始输出：\n{parsed['raw']}"
    
    response = f"您输入的车牌号对应的车场名称为：{parsed['park_name']}\n"
    
    if parsed["entry_lane"] and parsed["exit_lane"]:
        response += f"您输入的车牌号是从「{parsed['entry_lane']}」进入，准备从「{parsed['exit_lane']}」出去"
    elif parsed["entry_lane"]:
        response += f"入口车道：{parsed['entry_lane']}"
    elif parsed["exit_lane"]:
        response += f"出口车道：{parsed['exit_lane']}"
    
    return response


# --------------------------
# 4. 执行Linux命令（按你的要求拼接 ssh 前缀）
# --------------------------
def execute_linux_command(plate_number: str, user_text: str) -> str:
    if not plate_number:
        return "未提取到车牌"

    # ============================================================
    # 新增：优先处理"查询出场车道" / "查询车场名称" 请求
    # ============================================================
    if is_park_lane_query(user_text):
        print(f"【DEBUG】检测到车场/车道查询请求，车牌: {plate_number}")
        output = query_park_lane_info(plate_number)
        parsed = parse_park_lane_output(output)
        return format_park_lane_response(parsed, plate_number)

    filter_cmd = get_filter_command(user_text)

    # =========================
    # 1️⃣ 解析时间
    # =========================
    target_time = parse_time(user_text)

    # =========================
    # 2️⃣ 有时间 → 精准查日志
    # =========================
    if target_time:
        logs = list_remote_logs()

        if not logs:
            return "未找到日志文件"

        closest_log = find_closest_log(target_time, logs)
        filename = closest_log[0]

        print(f"【DEBUG】使用日志文件: {filename}")

        result = query_specific_log(filename, plate_number, filter_cmd)

        return result if result else "未查询到该车牌的相关日志记录"

    # =========================
    # 3️⃣ 无时间 → 标准流程
    # =========================

    # 3.1 查 app.log
    result = query_app_log(plate_number, filter_cmd)
    if result:
        return result

    # 3.2 查最近8小时
    logs = list_remote_logs()
    now = datetime.datetime.now()

    for filename, log_time in logs:
        if (now - log_time).total_seconds() <= 8 * 3600:
            result = query_specific_log(filename, plate_number, filter_cmd)
            if result:
                return result

    return "未查询到该车牌的相关日志记录"

def highlight_keywords(text: str) -> str:
    """高亮显示特定关键字：发送、欠费、白名单"""
    keywords = ["发送", "欠费", "白名单"]
    HIGHLIGHT = "\033[91m"  # 红色
    RESET = "\033[0m"
    result = text
    for kw in keywords:
        result = result.replace(kw, f"{HIGHLIGHT}{kw}{RESET}")
    return result
    

## 1️⃣ 图片映射表（你必须配置）

IMAGE_MAP = {
    "高新技术企业": "/root/dingding/images/gaoxin.jpg",
    "质量管理体系": "/root/dingding/images/quality.jpg",
    "环境管理体系": "/root/dingding/images/env.jpg",
    "职业健康": "/root/dingding/images/health.jpg",
    "信息安全": "/root/dingding/images/security.jpg",
    "法人": "/root/dingding/images/legal.jpg",
    "专利": "/root/dingding/images/patent.jpg",
    "合同": "/root/dingding/images/contract.jpg",
}

from docx.shared import Inches

def auto_fix_filename(filename: str, content: str) -> str:
    if filename:
        return filename

    # 尝试从内容中提取
    match = re.search(r'《(.+?)》', content)
    if match:
        return f"{match.group(1)}-投标文件.docx"

    return f"投标文件_{uuid.uuid4()}.docx"
    
def sanitize_filename(filename: str) -> str:
    filename = filename.strip()

    # 去掉非法字符
    filename = re.sub(r'[\\/:*?"<>|]', '_', filename)

    if not filename.endswith(".docx"):
        filename += ".docx"

    return filename
   
def generate_docx_file(content: str, filename: str = None):
    doc = Document()
    used_images = set()

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        # 标题加粗逻辑
        if line.startswith(("一、", "二、", "三、", "四、", "五、", "第")):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            doc.add_paragraph(line)
        # 插图逻辑
        for keyword, img_path in IMAGE_MAP.items():
            if keyword in line and keyword not in used_images:
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(4))
                    used_images.add(keyword)

    # 文件名处理（严格按SKILL.md规则）
    filename = auto_fix_filename(filename, content)  # 优先用户指定→提取招标文件名→fallback
    filename = sanitize_filename(filename)
    if "投标文件" not in filename:
        filename = filename.replace(".docx", "") + "-投标文件.docx"
    
    docs_dir = get_generated_docs_dir()
    docs_dir.mkdir(parents=True, exist_ok=True)
    filepath = str(docs_dir / filename)
    # 上面已经按行写过正文，这里不再重复追加整段 content，避免文档内容重复一份。
    doc.save(filepath)
    return filename, filepath
    
# ================== MCP Server ==================
app = Server("mysql_mcp_server")


# 注册工具列表接口，让客户端知道可调用的工具。
@app.list_tools()
async def list_tools() -> list[Tool]:
    return [
        Tool(
            name="execute-sql",
            description="执行SQL查询，支持分表路由、安全校验和自动限流",
            inputSchema={
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "要执行的SQL语句，必须包含 park.name 条件"
                    }
                },
                "required": ["query"]
            }
        ),
        Tool(
            name="execute-linux",
            inputSchema={
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "用户原始输入"
                    }
                },
                "required": ["query"]
            }
        ),
        Tool(
            name="generate-docx",
            description="生成Word文档、回写投标文件生成状态，并返回下载地址",
            inputSchema={
                "type": "object",
                "properties": {
                    "uid": {
                        "type": "string",
                        "description": "业务记录 uid，用于回写生成状态"
                    },
                    "tender_uid": {
                        "type": "string",
                        "description": "招标 tenderUid，用于回写生成状态"
                    },
                    "filename": {
                        "type": "string",
                        "description": "可选。生成的Word文件名（不传时会自动推断并补齐.docx）"
                    },
                    "content": {
                        "type": "string",
                        "description": "Word文档内容（markdown格式）"
                    }
                },
                "required": ["uid", "tender_uid", "content"]
            }
        )
    ]


@app.call_tool()
async def call_tool(name: str, arguments: dict):
    if name == "execute-sql":
        query = arguments.get("query")
        if not query:
            raise ValueError("Query required")

        try:
            validate_sql(query)
            query = rewrite_sql(query)
            #query = add_limit(query)

            logger.info(f"Final SQL: {query}")

            config = get_db_config()
            with connect(**config) as conn:
                with conn.cursor() as cursor:
                    cursor.execute(query)

                    if cursor.description:
                        columns = [c[0] for c in cursor.description]
                        rows = cursor.fetchall()
                        result = [",".join(map(str, r)) for r in rows]
                        return [
                            TextContent(
                                type="text",
                                text="\n".join([",".join(columns)] + result)
                            )
                        ]
                    else:
                        return [
                            TextContent(
                                type="text",
                                text=f"OK rows={cursor.rowcount}"
                            )
                        ]

        except Exception as e:
            logger.error(str(e))
            return [TextContent(type="text", text=f"Error: {str(e)}")]

    elif name == "execute-linux":
        user_text = arguments.get("query")  # 原始用户输入
        #user_text = request.params.arguments.get("query")  # 原始输入
        print(f"Before bad in:原始输入： {user_text}")
        if not user_text:
            raise ValueError("query required")
            
        for bad in ["前后时间", "zcat", "grep", "日志文件"]:
            user_text = user_text.replace(bad, "")

        plate = extract_license_plate(user_text)

        try:
            logger.info(f"Before execute linux command: 车牌：{plate}，原始输入：{user_text}")
            output = execute_linux_command(plate, user_text)  # 传入车牌 + 用户原始问题
            highlighted_output = highlight_keywords(output)
            return [
                TextContent(
                    type="text",
                    text=f"远程命令执行结果:\n\n{highlighted_output}"
                )
            ]

        except Exception as e:
            logger.error(str(e))
            return [TextContent(type="text", text=f"Error: {str(e)}")]

    elif name == "generate-docx":
        uid = str(arguments.get("uid", "")).strip()
        # 兼容 tender_uid 和 tenderUid 两种写法，避免 Agent 参数名大小写不一致导致无法回写。
        tender_uid = str(arguments.get("tender_uid") or arguments.get("tenderUid") or "").strip()
        content = arguments.get("content")
        filename = arguments.get("filename")

        try:
            if not uid:
                raise ValueError("uid required")
            if not tender_uid:
                raise ValueError("tender_uid required")
            if not content:
                raise ValueError("content required")

            filename, filepath = generate_docx_file(content, filename)
            download_url = f"{get_public_download_base_url()}/{quote(filename)}"

            status_resp = update_tender_generation_status(
                uid=uid,
                tender_uid=tender_uid,
                bid_doc_location_url=download_url,
                status=3,
                memo="标书已生成",
            )
            logger.info(
                "generate-docx success uid=%s tender_uid=%s file=%s status_resp=%s",
                uid,
                tender_uid,
                filepath,
                status_resp,
            )

            return [
                TextContent(
                    type="text",
                    text=f"""文件已生成：
{download_url}

状态已回写为生成成功。
如无法点击，请复制链接到浏览器打开。
"""
                )
            ]
        except Exception as e:
            logger.exception(
                "generate-docx failed uid=%s tender_uid=%s error=%s",
                uid,
                tender_uid,
                e,
            )
            fail_text = f"生成Word或回写状态失败：{e}"

            # 能拿到业务主键时，失败也在 MCP 里回写 status=4。
            if uid and tender_uid:
                try:
                    fail_resp = update_tender_generation_status(
                        uid=uid,
                        tender_uid=tender_uid,
                        status=4,
                        memo=fail_text,
                    )
                    logger.info(
                        "generate-docx failure status updated uid=%s tender_uid=%s status_resp=%s",
                        uid,
                        tender_uid,
                        fail_resp,
                    )
                except Exception as write_exc:
                    logger.exception(
                        "generate-docx failure status update failed uid=%s tender_uid=%s error=%s",
                        uid,
                        tender_uid,
                        write_exc,
                    )
                    fail_text = f"{fail_text}；失败状态回写也失败：{write_exc}"

            return [TextContent(type="text", text=f"Error: {fail_text}")]
    else:
        raise ValueError("Unknown tool")
    
# ================== HTTP服务 ==================
class StreamableHTTPASGIApp:
    def __init__(self, session_manager):
        self.session_manager = session_manager

    async def __call__(self,scope, receive, send):
        await self.session_manager.handle_request(scope, receive, send)


def create_http_app():
    session_manager = StreamableHTTPSessionManager(app=app, json_response=False)

    async def healthz(_):
        return JSONResponse({"status": "ok"})

    @asynccontextmanager
    async def lifespan(_):
        async with session_manager.run():
            yield

    from starlette.responses import FileResponse

    async def download(request):
        filename = request.path_params["filename"]
        filepath = str(get_generated_docs_dir() / filename)

        if not os.path.exists(filepath):
            return JSONResponse({"error": "file not found"}, status_code=404)

        return FileResponse(
            filepath,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
            content_disposition_type="attachment",
        )
# ================== 新增：skill-generator 路由 ==================
    async def skill_generator(request):
        filepath = "/root/skill/skill-generator-kimi.html"
        if not os.path.exists(filepath):
            return JSONResponse({"error": "skill generator file not found"}, status_code=404)
        return FileResponse(filepath, media_type="text/html")

    return Starlette(
        routes=[
            Route("/healthz", endpoint=healthz),
            Route("/mcp", endpoint=StreamableHTTPASGIApp(session_manager)),
            Route("/download/{filename}", endpoint=download),
            # ================== 在这里增加一行 ==================
            Route("/skill-generator", endpoint=skill_generator),
        ],
        lifespan=lifespan,
    )

async def main():
    http_app = create_http_app()

    config = uvicorn.Config(
        http_app,
        host="0.0.0.0",
        port=5418,
        log_level="info",
    )

    server = uvicorn.Server(config)
    await server.serve()


if __name__ == "__main__":
    asyncio.run(main())
